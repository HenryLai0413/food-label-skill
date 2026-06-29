#!/usr/bin/env python3
"""
食品標示 Word 文件產生器
字型：Noto Sans TC
檔名自動格式：產品標示_{品名}_{YYYYMMDD}.docx

使用方式：
  python generate_docx.py <input.json> [output_dir_or_path]

  若 output 為資料夾路徑（或省略），自動依品名與日期產生檔名。
  若 output 以 .docx 結尾，直接使用該路徑。

input_json 欄位：
  product_name        品名（純品名，不含素食括號）
  ingredients         成分字串（已依含量排序、複合原料已展開）
  nominal_weight_g    標示量（數字，公克），用於計算容許負誤差
  solid_content_g     固形物重量（公克，可選）
  expiry              有效日期（預設「標示於包裝上」）
  storage             保存方式
  origin              原產地（預設「台灣」）
  allergens           過敏原警語
  manufacturer        製造廠商名稱
  phone               電話
  address             地址
  factory_reg         工廠登記號碼
  serving_size        每份量（e.g. 250公克）
  servings_per_package 本包裝含幾份
  nutrition           營養數值 dict
  notes               備註（可選）
"""

import sys
import json
import os
import re
import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("請先安裝 python-docx：pip install python-docx")
    sys.exit(1)

FONT_NAME = "Noto Sans TC"


# ── 字型輔助 ─────────────────────────────────────────────────────────────────

def set_run_font(run, font_name=FONT_NAME, size_pt=None, bold=None, color=None):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_doc_default_font(doc, font_name=FONT_NAME):
    style = doc.styles["Normal"]
    style.font.name = font_name
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


# ── 淨重容許負誤差（依《定量包裝商品管理辦法》附表） ─────────────────────────

def calculate_tolerance(nominal_g: float) -> str:
    if nominal_g <= 0:
        return ""
    if nominal_g <= 50:
        return f"±{nominal_g * 0.09:.1f}公克（9%）"
    elif nominal_g <= 100:
        return "±4.5公克"
    elif nominal_g <= 200:
        return f"±{nominal_g * 0.045:.1f}公克（4.5%）"
    elif nominal_g <= 300:
        return "±9公克"
    elif nominal_g <= 500:
        return f"±{nominal_g * 0.03:.1f}公克（3%）"
    elif nominal_g <= 1000:
        return "±15公克"
    elif nominal_g <= 10000:
        return f"±{nominal_g * 0.015:.0f}公克（1.5%）"
    else:
        return "±150公克"


def format_net_weight(nominal_g: float, solid_g: float = None) -> str:
    tol = calculate_tolerance(nominal_g)
    base = f"{nominal_g:.0f}公克 {tol}"
    if solid_g and solid_g > 0:
        base += f"（固形物：{solid_g:.0f}公克）"
    return base


# ── 檔名產生 ─────────────────────────────────────────────────────────────────

def build_output_path(output_arg: str, product_name: str) -> str:
    """
    若 output_arg 為資料夾或不以 .docx 結尾，
    自動產生 產品標示_{品名}_{YYYYMMDD}.docx。
    """
    date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    # 清理品名：移除常見素食括號標記與檔名非法字元
    clean_name = product_name
    for pat in [r"（非素食）", r"（全素）", r"（純素）", r"（蛋素）",
                r"（奶素）", r"（蛋奶素）", r"\(.*?\)", r"（.*?）"]:
        clean_name = re.sub(pat, "", clean_name)
    # 移除 Windows 檔名非法字元
    clean_name = re.sub(r'[\\/:*?"<>|]', "", clean_name).strip()

    filename = f"產品標示_{clean_name}_{date_str}.docx"

    if not output_arg:
        return filename
    if output_arg.lower().endswith(".docx"):
        return output_arg
    # 視為資料夾
    os.makedirs(output_arg, exist_ok=True)
    return os.path.join(output_arg, filename)


# ── 表格輔助 ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_label_row(table, label: str, value: str, size_pt=10):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    lc.text = ""
    vc.text = ""
    lr = lc.paragraphs[0].add_run(label)
    set_run_font(lr, size_pt=size_pt, bold=True)
    vr = vc.paragraphs[0].add_run(value)
    set_run_font(vr, size_pt=size_pt)
    set_cell_bg(lc, "D9D9D9")
    return row


# ── 主產生函式 ───────────────────────────────────────────────────────────────

def generate_label_docx(data: dict, output_path: str):
    doc = Document()
    set_doc_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    product_name = data.get("product_name", "產品標示")
    vegetarian_type = data.get("vegetarian_type", "").strip()
    is_oem = bool(data.get("is_oem", False))

    # 品名顯示：素食產品在品名後加括號標註類型（台灣素食標示規範）
    display_name = f"{product_name}（{vegetarian_type}）" if vegetarian_type else product_name

    # 標題
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run(display_name)
    set_run_font(tr, size_pt=14, bold=True)

    doc.add_paragraph()

    # 淨重
    nominal_g = float(data.get("nominal_weight_g", 0) or 0)
    solid_g = float(data.get("solid_content_g", 0) or 0)
    net_weight_str = (
        format_net_weight(nominal_g, solid_g if solid_g > 0 else None)
        if nominal_g > 0
        else data.get("net_weight", "")
    )

    # 主標示表格（不含素食標示欄位）
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # OEM 產品將「製造廠商」改為「負責廠商」
    manufacturer_label = "負責廠商" if is_oem else "製造廠商"

    fields = [
        ("品名",              display_name),
        ("成分",              data.get("ingredients", "")),
        ("淨重",              net_weight_str),
        ("有效日期",          data.get("expiry", "標示於包裝上")),
        ("保存方式",          data.get("storage", "")),
        ("原產地",            data.get("origin", "台灣")),
        ("過敏原",            data.get("allergens", "")),
        (manufacturer_label,  data.get("manufacturer", "")),
        ("電話",              data.get("phone", "")),
        ("地址",              data.get("address", "")),
        ("工廠登記號碼",      data.get("factory_reg", "")),
    ]

    for label, value in fields:
        add_label_row(table, label, value)

    for row in table.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.4)

    doc.add_paragraph()

    # 營養標示
    nh_para = doc.add_paragraph()
    nh_run = nh_para.add_run("營養標示")
    set_run_font(nh_run, size_pt=11, bold=True)

    nutrition = data.get("nutrition", {})
    serving_size = data.get("serving_size", "")
    servings_per_pkg = data.get("servings_per_package", "")

    si_para = doc.add_paragraph()
    si_run = si_para.add_run(f"每份量：{serving_size}　　本包裝含：{servings_per_pkg}份")
    set_run_font(si_run, size_pt=10)

    nt = doc.add_table(rows=1, cols=3)
    nt.style = "Table Grid"
    nt.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["項目", "每份", "每100公克"]):
        cell = nt.rows[0].cells[i]
        cell.text = ""
        hr = cell.paragraphs[0].add_run(h)
        set_run_font(hr, size_pt=10, bold=True)
        set_cell_bg(cell, "D9D9D9")

    n = nutrition
    nut_rows = [
        ("熱量",       n.get("calories_per_serving", ""),      n.get("calories_per_100g", ""),      "大卡"),
        ("蛋白質",     n.get("protein_per_serving", ""),       n.get("protein_per_100g", ""),       "公克"),
        ("脂肪",       n.get("fat_per_serving", ""),           n.get("fat_per_100g", ""),           "公克"),
        ("　飽和脂肪", n.get("saturated_fat_per_serving", ""), n.get("saturated_fat_per_100g", ""), "公克"),
        ("　反式脂肪", n.get("trans_fat_per_serving", ""),     n.get("trans_fat_per_100g", ""),     "公克"),
        ("碳水化合物", n.get("carbs_per_serving", ""),         n.get("carbs_per_100g", ""),         "公克"),
        ("　糖",       n.get("sugar_per_serving", ""),         n.get("sugar_per_100g", ""),         "公克"),
        ("鈉",         n.get("sodium_per_serving", ""),        n.get("sodium_per_100g", ""),        "毫克"),
    ]

    for item_name, per_serving, per_100g, unit in nut_rows:
        row = nt.add_row()
        for cell, txt in zip(row.cells, [
            item_name,
            f"{per_serving} {unit}" if per_serving else "",
            f"{per_100g} {unit}" if per_100g else "",
        ]):
            cell.text = ""
            r = cell.paragraphs[0].add_run(txt)
            set_run_font(r, size_pt=10)

    # 備註
    notes = data.get("notes", "")
    if notes:
        doc.add_paragraph()
        np_ = doc.add_paragraph()
        nr = np_.add_run(f"備註：{notes}")
        set_run_font(nr, size_pt=9, color=RGBColor(0x80, 0x80, 0x80))

    # 法規聲明
    doc.add_paragraph()
    dp = doc.add_paragraph()
    dr = dp.add_run(
        "※ 本標示草稿依《食品安全衛生管理法》第22條產出，"
        "淨重容許負誤差依《定量包裝商品管理辦法》附表計算；"
        "營養值參考台灣食品成分資料庫2025版（衛生福利部食藥署），請以原料規格書最終校正。"
    )
    set_run_font(dr, size_pt=8, color=RGBColor(0x80, 0x80, 0x80))

    doc.save(output_path)
    print(f"產品標示文件已產出：{output_path}")
    print(f"字型：{FONT_NAME}  |  淨重容許負誤差：{calculate_tolerance(nominal_g) if nominal_g > 0 else 'N/A'}")


DEFAULT_OUTPUT_DIR = r"C:\Users\hchen\Desktop\SynologyDrive\03【存】產品 (照片, 條碼, 檢驗報告, 目錄)"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python generate_docx.py <input.json> [output_dir_or_.docx_path]")
        sys.exit(1)

    input_arg = sys.argv[1]
    output_arg = sys.argv[2] if len(sys.argv) >= 3 else DEFAULT_OUTPUT_DIR

    if os.path.exists(input_arg):
        with open(input_arg, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.loads(input_arg)

    final_path = build_output_path(output_arg, data.get("product_name", "產品"))
    generate_label_docx(data, final_path)
